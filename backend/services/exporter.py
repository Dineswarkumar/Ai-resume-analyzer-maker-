import os
import uuid
from datetime import datetime
from models import ResumeData
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

EXPORTS_DIR = "exports"
os.makedirs(EXPORTS_DIR, exist_ok=True)


# ─── PDF Exporter ────────────────────────────────────────────────────────────

class ATSProTemplate:
    """Clean, single-column ATS-optimized layout."""

    PRIMARY = colors.HexColor("#1A1A2E")
    ACCENT = colors.HexColor("#7C3AED")
    GRAY = colors.HexColor("#6B7280")
    LIGHT = colors.HexColor("#F3F4F6")

    def build_styles(self):
        styles = {}
        styles["name"] = ParagraphStyle(
            "name", fontSize=22, fontName="Helvetica-Bold",
            textColor=self.PRIMARY, alignment=TA_LEFT, spaceAfter=2
        )
        styles["contact"] = ParagraphStyle(
            "contact", fontSize=9, fontName="Helvetica",
            textColor=self.GRAY, spaceAfter=10
        )
        styles["section_heading"] = ParagraphStyle(
            "section_heading", fontSize=11, fontName="Helvetica-Bold",
            textColor=self.ACCENT, spaceBefore=10, spaceAfter=4
        )
        styles["body"] = ParagraphStyle(
            "body", fontSize=9.5, fontName="Helvetica",
            textColor=self.PRIMARY, spaceAfter=3, leading=14
        )
        styles["bullet"] = ParagraphStyle(
            "bullet", fontSize=9, fontName="Helvetica",
            textColor=self.PRIMARY, spaceAfter=2, leftIndent=12,
            bulletIndent=4, leading=13
        )
        styles["sub_heading"] = ParagraphStyle(
            "sub_heading", fontSize=10, fontName="Helvetica-Bold",
            textColor=self.PRIMARY, spaceAfter=1
        )
        styles["sub_info"] = ParagraphStyle(
            "sub_info", fontSize=9, fontName="Helvetica-Oblique",
            textColor=self.GRAY, spaceAfter=2
        )
        return styles

    def build(self, resume: ResumeData, filepath: str):
        doc = SimpleDocTemplate(
            filepath, pagesize=A4,
            leftMargin=1.5*cm, rightMargin=1.5*cm,
            topMargin=1.5*cm, bottomMargin=1.5*cm
        )
        s = self.build_styles()
        story = []

        # ── Name & Contact ──
        story.append(Paragraph(resume.contact.name or "Your Name", s["name"]))
        contact_parts = filter(None, [
            resume.contact.email, resume.contact.phone,
            resume.contact.linkedin, resume.contact.github,
            resume.contact.location
        ])
        story.append(Paragraph(" | ".join(contact_parts), s["contact"]))
        story.append(HRFlowable(width="100%", thickness=1.5, color=self.ACCENT))

        # ── Summary ──
        if resume.summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", s["section_heading"]))
            story.append(Paragraph(resume.summary, s["body"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=self.LIGHT))

        # ── Skills ──
        if resume.skills:
            story.append(Paragraph("SKILLS & TECHNOLOGIES", s["section_heading"]))
            story.append(Paragraph(", ".join(resume.skills), s["body"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=self.LIGHT))

        # ── Experience ──
        if resume.experience:
            story.append(Paragraph("WORK EXPERIENCE", s["section_heading"]))
            for exp in resume.experience:
                story.append(Paragraph(f"{exp.title} — {exp.company}", s["sub_heading"]))
                if exp.duration:
                    story.append(Paragraph(exp.duration, s["sub_info"]))
                for ach in exp.achievements:
                    story.append(Paragraph(f"• {ach}", s["bullet"]))
                if exp.description and not exp.achievements:
                    story.append(Paragraph(exp.description, s["body"]))
                story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=self.LIGHT))

        # ── Projects ──
        if resume.projects:
            story.append(Paragraph("PROJECTS", s["section_heading"]))
            for proj in resume.projects:
                techs = ", ".join(proj.technologies) if proj.technologies else ""
                title = f"<b>{proj.name}</b>"
                if techs:
                    title += f" <font color='#7C3AED'>| {techs}</font>"
                story.append(Paragraph(title, s["body"]))
                if proj.description:
                    story.append(Paragraph(proj.description, s["bullet"]))
                if proj.github or proj.url:
                    link = proj.github or proj.url
                    story.append(Paragraph(f"🔗 {link}", s["bullet"]))
                story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=self.LIGHT))

        # ── Education ──
        if resume.education:
            story.append(Paragraph("EDUCATION", s["section_heading"]))
            for edu in resume.education:
                story.append(Paragraph(edu.degree, s["sub_heading"]))
                info_parts = filter(None, [edu.institution, edu.year, f"GPA: {edu.gpa}" if edu.gpa else ""])
                story.append(Paragraph(" | ".join(info_parts), s["sub_info"]))
                story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=self.LIGHT))

        # ── Certifications ──
        if resume.certifications:
            story.append(Paragraph("CERTIFICATIONS", s["section_heading"]))
            for cert in resume.certifications:
                line = cert.name
                if cert.issuer:
                    line += f" — {cert.issuer}"
                if cert.year:
                    line += f" ({cert.year})"
                story.append(Paragraph(f"• {line}", s["bullet"]))

        doc.build(story)


class ModernCreativeTemplate(ATSProTemplate):
    """Two-column modern layout with sidebar."""

    SIDEBAR_BG = colors.HexColor("#1A1A2E")
    SIDEBAR_TEXT = colors.white
    ACCENT = colors.HexColor("#06D6A0")

    def build(self, resume: ResumeData, filepath: str):
        # For simplicity, use the ATS layout with different colors
        # A full two-column layout needs canvas overrides — keeping clean for now
        super().build(resume, filepath)


class AcademicTemplate(ATSProTemplate):
    """Formal academic/research layout."""

    PRIMARY = colors.HexColor("#1C1C1C")
    ACCENT = colors.HexColor("#C0392B")

    def build(self, resume: ResumeData, filepath: str):
        super().build(resume, filepath)


TEMPLATE_MAP = {
    "ats_pro": ATSProTemplate,
    "modern_creative": ModernCreativeTemplate,
    "academic": AcademicTemplate,
}


def export_to_pdf(resume: ResumeData, template_id: str = "ats_pro") -> str:
    filename = f"resume_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d')}.pdf"
    filepath = os.path.join(EXPORTS_DIR, filename)
    TemplateClass = TEMPLATE_MAP.get(template_id, ATSProTemplate)
    builder = TemplateClass()
    builder.build(resume, filepath)
    return filepath


def export_to_docx(resume: ResumeData) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Style helpers
    def add_heading(text, level=1, color=(124, 58, 237)):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
        return h

    def add_paragraph(text, bold=False, italic=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume.contact.name or "Your Name")
    name_run.font.size = Pt(22)
    name_run.bold = True
    name_run.font.color.rgb = RGBColor(26, 26, 46)

    # Contact
    contact_parts = filter(None, [
        resume.contact.email, resume.contact.phone,
        resume.contact.linkedin, resume.contact.github
    ])
    add_paragraph(" | ".join(contact_parts), size=9)

    doc.add_paragraph("─" * 80)

    if resume.summary:
        add_heading("PROFESSIONAL SUMMARY", 2)
        add_paragraph(resume.summary)

    if resume.skills:
        add_heading("SKILLS & TECHNOLOGIES", 2)
        add_paragraph(", ".join(resume.skills))

    if resume.experience:
        add_heading("WORK EXPERIENCE", 2)
        for exp in resume.experience:
            add_paragraph(f"{exp.title} — {exp.company}", bold=True)
            if exp.duration:
                add_paragraph(exp.duration, italic=True, size=9)
            for ach in exp.achievements:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(ach).font.size = Pt(9)

    if resume.projects:
        add_heading("PROJECTS", 2)
        for proj in resume.projects:
            techs = f" ({', '.join(proj.technologies)})" if proj.technologies else ""
            add_paragraph(f"{proj.name}{techs}", bold=True)
            if proj.description:
                add_paragraph(proj.description, size=9)

    if resume.education:
        add_heading("EDUCATION", 2)
        for edu in resume.education:
            add_paragraph(edu.degree, bold=True)
            info = " | ".join(filter(None, [edu.institution, edu.year]))
            add_paragraph(info, italic=True, size=9)

    if resume.certifications:
        add_heading("CERTIFICATIONS", 2)
        for cert in resume.certifications:
            line = cert.name
            if cert.year:
                line += f" ({cert.year})"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line).font.size = Pt(9)

    filename = f"resume_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(EXPORTS_DIR, filename)
    doc.save(filepath)
    return filepath
