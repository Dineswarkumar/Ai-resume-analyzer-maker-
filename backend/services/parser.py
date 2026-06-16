import re
import io
from typing import Tuple
from models import (
    ResumeData, ContactInfo, Education, Experience, Project, Certification
)

# ──────────────────────────────────────────────────────
#  Raw text extractors per file type
# ──────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        return text
    except Exception as e:
        raise ValueError(f"PDF parse error: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"DOCX parse error: {e}")


def extract_text_from_excel(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        raise ValueError(f"Excel parse error: {e}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        raise ValueError(f"TXT parse error: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return extract_text_from_excel(file_bytes)
    elif ext == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        # Try as plain text fallback
        return extract_text_from_txt(file_bytes)


# ──────────────────────────────────────────────────────
#  NLP / Regex entity extraction
# ──────────────────────────────────────────────────────

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\s\-().]{7,15}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)
GITHUB_RE = re.compile(r"github\.com/[\w\-]+", re.IGNORECASE)
URL_RE = re.compile(r"https?://[^\s]+")

SECTION_HEADERS = {
    "experience": ["experience", "work experience", "employment", "professional experience", "internship"],
    "education": ["education", "academic", "qualification", "degree"],
    "skills": ["skills", "technical skills", "core competencies", "technologies", "expertise"],
    "projects": ["projects", "personal projects", "academic projects", "portfolio"],
    "certifications": ["certifications", "certificates", "licenses", "achievements", "awards"],
    "summary": ["summary", "objective", "profile", "about me", "overview"],
}

COMMON_SKILLS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "kotlin", "swift",
    "react", "angular", "vue", "node.js", "express", "django", "flask", "fastapi", "spring",
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "docker", "kubernetes", "aws", "azure", "gcp", "terraform", "ci/cd",
    "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch",
    "pandas", "numpy", "scikit-learn", "matplotlib", "seaborn",
    "git", "linux", "rest api", "graphql", "microservices", "agile", "scrum",
    "html", "css", "tailwind", "bootstrap", "figma", "photoshop",
    "excel", "power bi", "tableau", "r", "matlab", "hadoop", "spark",
    # Extra common skills (often in student resumes)
    "php", "c programming", "c language", "basic ml", "data analysis",
    "github", "vs code", "linux", "shell scripting", "flask", "opencv",
]

DEGREE_KEYWORDS = ["b.tech", "b.e.", "btech", "b.sc", "bsc", "m.tech", "mtech", "m.sc",
                   "mba", "phd", "ph.d", "bachelor", "master", "associate", "diploma",
                   "10th", "11th", "12th", "ssc", "hsc", "intermediate", "matriculation"]

ACTION_VERBS = ["developed", "built", "designed", "implemented", "optimized", "led",
                "managed", "created", "deployed", "automated", "improved", "achieved",
                "increased", "reduced", "collaborated", "analyzed", "researched"]


def detect_sections(text: str) -> dict:
    """Split resume text into labelled sections."""
    lines = text.split("\n")
    sections = {k: [] for k in SECTION_HEADERS}
    sections["other"] = []
    current = "other"

    for line in lines:
        line_lower = line.lower().strip()
        matched = False
        for section, keywords in SECTION_HEADERS.items():
            if any(line_lower == kw or line_lower.startswith(kw) for kw in keywords):
                current = section
                matched = True
                break
        if not matched:
            sections[current].append(line)

    return {k: "\n".join(v) for k, v in sections.items()}


def extract_contact(text: str) -> ContactInfo:
    email = EMAIL_RE.search(text)
    phone = PHONE_RE.search(text)
    github = GITHUB_RE.search(text)

    # LinkedIn: handle URLs split across lines (e.g. "linkedin.com/in/dineswara-\nkumar-t13901s")
    # Join all text to one line for URL matching
    oneline = " ".join(text.split())
    linkedin = LINKEDIN_RE.search(oneline)

    # Name heuristic: collect first 1-2 non-empty all-caps or title-case lines
    name = ""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    name_parts = []
    for line in lines[:6]:
        # Skip lines that look like contact info
        if EMAIL_RE.search(line) or PHONE_RE.search(line):
            continue
        if line.startswith(":") or line.startswith("www") or line.startswith("linkedin"):
            continue
        # Accept all-caps short words or title-case words as name parts
        words = line.split()
        if 1 <= len(words) <= 4 and all(w.replace(',','').replace('.','').isalpha() for w in words):
            name_parts.append(line.strip())
            if len(" ".join(name_parts).split()) >= 2:
                break
    name = " ".join(name_parts).strip()
    # Title-case the name for clean display
    name = name.title() if name.isupper() else name

    # Location: look for city/country patterns
    location = ""
    loc_match = re.search(
        r"(bangalore|hyderabad|mumbai|delhi|chennai|pune|kolkata|nagpur|india|\w+,\s*india)",
        text, re.IGNORECASE
    )
    if loc_match:
        location = loc_match.group().title()

    return ContactInfo(
        name=name,
        email=email.group() if email else "",
        phone=phone.group().strip() if phone else "",
        linkedin="https://" + linkedin.group() if linkedin else "",
        github="https://" + github.group() if github else "",
        location=location,
    )


# Words that appear in resume sections but are NOT skills
NON_SKILL_WORDS = {
    "hobbies", "languages", "english", "hindi", "telugu", "tamil", "kannada", "marathi",
    "declaration", "signature", "place", "date", "nationality", "indian", "male", "female",
    "reading", "travelling", "dancing", "cooking", "photography", "gaming", "sports",
    "learning", "novels", "channels", "websites", "music", "hobbies", "interests",
    "reference", "references", "available", "upon request", "skills", "technical",
}

def extract_skills(skills_text: str, full_text: str) -> list:
    combined = (skills_text + " " + full_text).lower()
    found = []
    for skill in COMMON_SKILLS:
        if skill.lower() in combined and skill not in found:
            found.append(skill)

    # Also extract bullet-list items from skills section (each line is a skill)
    if skills_text:
        for line in skills_text.split("\n"):
            line = line.strip().lstrip("- *\u2022\u25ba\u25b8").strip()
            # Accept lines that are 2-40 chars
            if 2 <= len(line) <= 40 and line:
                # Skip if it's a non-skill word
                if line.lower().strip() in NON_SKILL_WORDS:
                    continue
                # Skip section header-like words (ALL CAPS single word)
                if line.isupper() and len(line.split()) == 1:
                    continue
                # Skip lines with special chars that indicate contact info
                if any(c in line for c in ['@', 'http', 'www', ':', '/']):  
                    continue
                canonical = line.lower()
                if canonical not in [s.lower() for s in found]:
                    found.append(line.title() if line.isupper() else line)
    return found


def extract_education(edu_text: str) -> list:
    entries = []
    lines = [l.strip() for l in edu_text.split("\n") if l.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        line_lower = line.lower()
        if any(kw in line_lower for kw in DEGREE_KEYWORDS):
            edu = Education()
            edu.degree = line
            # Look ahead for institution and year
            if i + 1 < len(lines):
                edu.institution = lines[i + 1]
            # Look for year pattern
            year_match = re.search(r"\b(19|20)\d{2}\b", " ".join(lines[max(0, i-1):i+4]))
            if year_match:
                edu.year = year_match.group()
            gpa_match = re.search(r"gpa[:\s]*([\d.]+)", " ".join(lines[i:i+3]), re.IGNORECASE)
            if gpa_match:
                edu.gpa = gpa_match.group(1)
            entries.append(edu)
            i += 2
        else:
            i += 1
    return entries if entries else []


def extract_experience(exp_text: str) -> list:
    entries = []
    blocks = re.split(r"\n{2,}", exp_text)
    for block in blocks:
        if not block.strip():
            continue
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not lines:
            continue
        exp = Experience()
        exp.title = lines[0] if lines else ""
        exp.company = lines[1] if len(lines) > 1 else ""
        # Duration
        duration_match = re.search(
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|"
            r"March|April|June|July|August|September|October|November|December)[\s,]*"
            r"(\d{4})\s*[-–]\s*(Present|\w+[\s,]*\d{4})",
            block, re.IGNORECASE
        )
        if duration_match:
            exp.duration = duration_match.group()
        # Achievements = bullet lines with action verbs
        exp.achievements = [
            l.lstrip("•-*").strip()
            for l in lines[2:]
            if any(v in l.lower() for v in ACTION_VERBS)
        ]
        exp.description = "\n".join(lines[2:])
        if exp.title:
            entries.append(exp)
    return entries


def extract_projects(proj_text: str) -> list:
    entries = []
    blocks = re.split(r"\n{2,}", proj_text)
    for block in blocks:
        if not block.strip():
            continue
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not lines:
            continue
        proj = Project()
        proj.name = lines[0]
        proj.description = " ".join(lines[1:3]) if len(lines) > 1 else ""
        # Extract tech stack from parentheses or "Technologies: ..."
        tech_match = re.search(r"\(([^)]+)\)", block)
        if tech_match:
            proj.technologies = [t.strip() for t in tech_match.group(1).split(",")]
        # URL
        url_match = URL_RE.search(block)
        if url_match:
            url = url_match.group()
            if "github" in url:
                proj.github = url
            else:
                proj.url = url
        if proj.name:
            entries.append(proj)
    return entries


def extract_certifications(cert_text: str) -> list:
    entries = []
    lines = [l.strip() for l in cert_text.split("\n") if l.strip()]
    for line in lines:
        if len(line) < 5:
            continue
        cert = Certification()
        cert.name = line
        year_match = re.search(r"\b(20\d{2})\b", line)
        if year_match:
            cert.year = year_match.group()
        entries.append(cert)
    return entries[:10]  # cap at 10


# ──────────────────────────────────────────────────────
#  Main parse function
# ──────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> Tuple[ResumeData, str, float]:
    """
    Returns (ResumeData, raw_text, confidence_score 0-1)
    """
    raw_text = extract_text(file_bytes, filename)
    sections = detect_sections(raw_text)

    contact = extract_contact(raw_text)
    skills = extract_skills(sections.get("skills", ""), raw_text)
    education = extract_education(sections.get("education", ""))
    experience = extract_experience(sections.get("experience", ""))
    projects = extract_projects(sections.get("projects", ""))
    certifications = extract_certifications(sections.get("certifications", ""))

    # Summary: use detected summary section OR the "other" section's first large paragraph
    summary = sections.get("summary", "").strip()
    if not summary:
        other = sections.get("other", "")
        # Find the first paragraph with 20+ words — likely the summary
        paragraphs = [p.strip() for p in re.split(r"\n{2,}", other) if p.strip()]
        for para in paragraphs:
            word_count = len(para.split())
            if word_count >= 20:
                summary = para[:500]
                break

    # Languages: extract from LANGUAGES section if present
    languages = []
    lang_text = sections.get("other", "")
    lang_match = re.search(
        r"LANGUAGES?\s*\n(.*?)(?:\n\n|$)", lang_text, re.IGNORECASE | re.DOTALL
    )
    if lang_match:
        lang_lines = [l.strip() for l in lang_match.group(1).split("\n") if l.strip()]
        languages = [l for l in lang_lines if 2 <= len(l) <= 30 and l.isalpha()]

    resume_data = ResumeData(
        contact=contact,
        summary=summary[:500],
        education=education,
        experience=experience,
        skills=skills,
        projects=projects,
        certifications=certifications,
        languages=languages,
    )

    # Confidence: proportion of fields filled
    filled = sum([
        bool(contact.name), bool(contact.email), bool(contact.phone),
        bool(skills), bool(education), bool(experience or projects)
    ])
    confidence = filled / 6.0

    return resume_data, raw_text, confidence
